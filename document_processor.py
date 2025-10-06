"""
Document Processor - Handles DOCX and PDF text extraction
"""

import io
from typing import Union
import re


class DocumentProcessor:
    """Extracts text content from DOCX and PDF files"""
    
    def extract_text(self, content: bytes, content_type: str) -> str:
        """
        Extract text from document content
        
        Args:
            content: Raw file bytes
            content_type: MIME type of the document
            
        Returns:
            Extracted text content
        """
        if content_type == "application/pdf":
            return self._extract_from_pdf(content)
        elif content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self._extract_from_docx(content)
        else:
            raise ValueError(f"Unsupported content type: {content_type}")
    
    def _extract_from_pdf(self, content: bytes) -> str:
        """Extract text from PDF using PyPDF2"""
        try:
            import PyPDF2
            
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = []
            for page in pdf_reader.pages:
                text.append(page.extract_text())
            
            return "\n\n".join(text)
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF processing. Install with: pip install PyPDF2")
    
    def _extract_from_docx(self, content: bytes) -> str:
        """Extract text from DOCX using python-docx"""
        try:
            from docx import Document
            
            doc_file = io.BytesIO(content)
            doc = Document(doc_file)
            
            text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text.append(row_text)
            
            return "\n\n".join(text)
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
    
    def clean_text(self, text: str) -> str:
        """
        Clean extracted text - remove excessive whitespace, normalize line breaks
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Normalize paragraph breaks
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()